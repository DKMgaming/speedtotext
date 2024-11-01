import os
import streamlit as st
import speech_recognition as sr
from docx import Document
import librosa
import soundfile as sf

# Thiết lập thư mục tải lên
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Hàm chuyển đổi âm thanh sang văn bản
def convert_audio_to_text(audio_path):
    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_path) as source:
        audio = recognizer.record(source)
    try:
        text = recognizer.recognize_google(audio, language='vi-VN')
    except sr.UnknownValueError:
        text = "Không thể nhận diện được giọng nói"
    except sr.RequestError:
        text = "Lỗi khi yêu cầu dịch vụ nhận diện giọng nói"
    return text

# Hàm tạo tài liệu Word từ văn bản
def create_word_document(text, output_path):
    doc = Document()
    doc.add_heading('Transcript', level=1)
    doc.add_paragraph(text)
    doc.save(output_path)

# Giao diện Streamlit
st.title("Chuyển đổi âm thanh thành văn bản")
st.write("Tải lên file âm thanh để nhận diện và tạo tài liệu Word.")

# Tải file
uploaded_file = st.file_uploader("Chọn file âm thanh", type=["wav", "mp3", "m4a"])

if uploaded_file is not None:
    # Lưu file tải lên vào thư mục
    audio_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
    with open(audio_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Chuyển đổi file MP3 sang WAV nếu cần
    if uploaded_file.type == "audio/mpeg":
        wav_audio_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name.replace(".mp3", ".wav"))
        # Sử dụng librosa để chuyển đổi MP3 sang WAV
        audio_data, sr = librosa.load(audio_path, sr=None)  # Đọc file âm thanh
        sf.write(wav_audio_path, audio_data, sr)  # Lưu thành file WAV
        audio_path = wav_audio_path  # Cập nhật đường dẫn đến file WAV

    # Chuyển đổi âm thanh sang văn bản
    text = convert_audio_to_text(audio_path)

    # Tạo tài liệu Word
    output_path = os.path.join(UPLOAD_FOLDER, 'transcript.docx')
    create_word_document(text, output_path)

    # Tải xuống tài liệu Word
    with open(output_path, "rb") as f:
        st.download_button("Tải xuống tài liệu Word", f, file_name='transcript.docx')
